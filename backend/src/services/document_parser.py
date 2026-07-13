"""Document parsing helpers for uploaded files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any


class DocumentParseError(ValueError):
    """Raised when a document cannot produce searchable text."""


def parse_document(filename: str, content_type: str, payload: bytes, config: Any | None = None) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix in {"txt", "md"}:
        return _parse_text(payload), _content_type_for_suffix(suffix, content_type)
    if suffix == "pdf":
        return _parse_pdf(payload, config), "application/pdf"
    if suffix == "docx":
        return _parse_docx(payload), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise DocumentParseError("Only .txt, .md, .pdf and .docx documents are supported")


def _parse_text(payload: bytes) -> str:
    try:
        text = payload.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("Text documents must use UTF-8 encoding") from exc
    if not text.strip():
        raise DocumentParseError("Document has no searchable text")
    return text


def _parse_pdf(payload: bytes, config: Any | None = None) -> str:
    text = _extract_pdf_text(payload)
    if text.strip():
        return text
    if not getattr(config, "pdf_ocr_enabled", False):
        raise DocumentParseError("PDF has no searchable text; OCR is disabled")
    return _ocr_pdf(payload, config)


def _extract_pdf_text(payload: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - only hit when runtime deps are missing
        raise DocumentParseError("Missing pypdf dependency; cannot parse PDF") from exc

    try:
        reader = PdfReader(BytesIO(payload))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError("PDF parse failed") from exc
    return "\n\n".join(page for page in pages if page)


def _ocr_pdf(payload: bytes, config: Any) -> str:
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError as exc:  # pragma: no cover - depends on optional runtime deps
        raise DocumentParseError("PDF OCR requires optional dependencies: pdf2image, pytesseract and pillow") from exc

    tesseract_cmd = getattr(config, "tesseract_cmd", None)
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    try:
        images = convert_from_bytes(
            payload,
            dpi=max(72, int(getattr(config, "pdf_ocr_dpi", 200))),
            first_page=1,
            last_page=max(1, int(getattr(config, "pdf_ocr_max_pages", 20))),
            poppler_path=getattr(config, "poppler_path", None) or None,
        )
        parts = [
            pytesseract.image_to_string(image, lang=getattr(config, "pdf_ocr_language", "chi_sim+eng")).strip()
            for image in images
        ]
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError("PDF OCR failed; check Tesseract and Poppler installation") from exc

    text = "\n\n".join(part for part in parts if part)
    if not text.strip():
        raise DocumentParseError("PDF OCR produced no searchable text")
    return text


def _parse_docx(payload: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - only hit when runtime deps are missing
        raise DocumentParseError("Missing python-docx dependency; cannot parse DOCX") from exc

    try:
        document = Document(BytesIO(payload))
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError("DOCX parse failed") from exc

    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    if not text.strip():
        raise DocumentParseError("DOCX has no searchable text")
    return text


def _content_type_for_suffix(suffix: str, content_type: str) -> str:
    if suffix == "md":
        return content_type or "text/markdown"
    return content_type or "text/plain"
